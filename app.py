import base64
import re
import io
import csv
import os
from shutil import copyfile
import zipfile
from flask import Flask, json, jsonify, render_template, request, send_file, send_from_directory
from werkzeug.utils import secure_filename
from fpdf import FPDF
from PIL import Image
from reportlab.pdfgen import canvas
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import logging
import pandas as pd
from docx import Document

temp_folder = "output"

# Настройка логирования
logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)

# Конфигурация приложения
UPLOAD_FOLDER = "uploads"
OUTPUT_FOLDER = "output"
FONT_PATH = "static/fonts/DejaVuSans.ttf"
TEMPLATE_IMAGE_PATH = "static/images/template1_ext.jpg"
ALLOWED_EXTENSIONS = {'csv'}
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

# Регистрация шрифта с поддержкой Unicode
pdfmetrics.registerFont(TTFont("DejaVu", FONT_PATH))

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

def extract_tags(doc):
    """Извлекает все уникальные теги из документа, учитывая разбиение на runs"""
    tags = set()
    pattern = re.compile(r'(%\w+)')
    
    # Обработка параграфов
    for paragraph in doc.paragraphs:
        full_text = ''.join(run.text for run in paragraph.runs)
        tags.update(pattern.findall(full_text))
    
    # Обработка таблиц
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    full_text = ''.join(run.text for run in paragraph.runs)
                    tags.update(pattern.findall(full_text))
    
    # Обработка колонтитулов
    for section in doc.sections:
        # Header
        for paragraph in section.header.paragraphs:
            full_text = ''.join(run.text for run in paragraph.runs)
            tags.update(pattern.findall(full_text))
        # Footer
        for paragraph in section.footer.paragraphs:
            full_text = ''.join(run.text for run in paragraph.runs)
            tags.update(pattern.findall(full_text))
    
    return tags

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def generate_certificate(name):
    try:
        output_filename = os.path.join(OUTPUT_FOLDER, f"{name.replace(' ', '_')}.pdf")

        pdf = FPDF('P', 'mm', 'A4')
        pdf.add_page()
        pdf.add_font('DejaVu', '', FONT_PATH, uni=True)
        pdf.set_font('DejaVu', '', 22)

        pdf.image(TEMPLATE_IMAGE_PATH, x=0, y=0, w=210, h=297, type='JPG')
        pdf.ln(90)
        pdf.cell(w=190, h=40, align='C', txt=name, border=False)

        pdf.output(output_filename)
        logger.info(f"Сертификат создан: {output_filename}")
    except Exception as e:
        logger.error(f"Ошибка при создании сертификата для {name}: {e}")

def process_csv_file(file):
    records = []
    with open(file, encoding='utf-8-sig') as csvfile:
        reader = csv.reader(csvfile)
        for row in reader:
            cleaned_row = re.sub(r'\s+', ' ', row[0]).strip()
            if cleaned_row:
                records.append(cleaned_row)
    return records

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/template1')
def template1():
    return render_template('template1.html')

@app.route('/your_temp')
def your_temp():
    return render_template('your_temp.html')

@app.route('/generate-pdf', methods=['POST'])
def generate_pdf():
    try:
        full_name = request.form['fullName']

        pdf = FPDF('P', 'mm', 'A4')
        pdf.add_page()
        pdf.add_font('DejaVu', '', FONT_PATH, uni=True)
        pdf.set_font('DejaVu', '', 22)

        pdf.image(TEMPLATE_IMAGE_PATH, x=0, y=0, w=210, h=297, type='JPG')
        pdf.ln(90)
        pdf.cell(w=180, h=40, align='C', txt=full_name, border=False)

        output_filename = os.path.join(OUTPUT_FOLDER, "certificate.pdf")
        pdf.output(output_filename)

        with open(output_filename, "rb") as file:
            data = file.read()

        return send_file(
            io.BytesIO(data),
            mimetype='application/pdf',
            as_attachment=True,
            download_name=f'{full_name}_certificate.pdf'
        )
    except Exception as e:
        logger.error(f"Ошибка при генерации PDF: {e}")
        return jsonify({"error": "Ошибка при генерации PDF."}), 500

@app.route('/process-csv', methods=['POST'])
def process_csv():
    if 'csvfile' not in request.files:
        return jsonify({'error': 'CSV-файл не был загружен.'}), 400

    file = request.files['csvfile']
    if file.filename == '' or not allowed_file(file.filename):
        return jsonify({'error': 'Неверный формат файла.'}), 400

    try:
        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(filepath)

        records = process_csv_file(filepath)
        for record in records:
            generate_certificate(record)

        return jsonify({'count': len(records)})
    except Exception as e:
        logger.error(f"Ошибка при обработке CSV: {e}")
        return jsonify({'error': 'Ошибка при обработке CSV.'}), 500

@app.route('/save-template', methods=['POST'])
def save_template():
    try:
        data = request.json
        image_data = data.get('imageData')

        if not image_data:
            raise ValueError("Нет данных изображения")

        header, encoded = image_data.split(",", 1)
        image_bytes = io.BytesIO(base64.b64decode(encoded))
        img = Image.open(image_bytes)

        if img.mode == 'RGBA':
            img = img.convert('RGB')

        temp_image_path = os.path.join(OUTPUT_FOLDER, "temp_image.jpg")
        img.save(temp_image_path, format="JPEG")

        output_pdf_path = os.path.join(OUTPUT_FOLDER, "template_with_text.pdf")
        pdf_canvas = canvas.Canvas(output_pdf_path, pagesize=(img.width, img.height))

        pdf_canvas.drawImage(temp_image_path, 0, 0, width=img.width, height=img.height)
        pdf_canvas.save()

        os.remove(temp_image_path)

        return jsonify({"success": True, "downloadUrl": f"/download/{output_pdf_path}"})
    except Exception as e:
        logger.error(f"Ошибка при сохранении шаблона: {e}")
        return jsonify({"success": False, "error": str(e)})

@app.route('/upload-csv', methods=['POST'])
def upload_csv():
    # Проверяем, был ли загружен файл
    if 'csvFile' not in request.files:
        return jsonify(success=False, error="Файл CSV не был загружен."), 400

    file = request.files['csvFile']

    # Проверяем формат файла
    if not file.filename.endswith('.csv'):
        return jsonify(success=False, error="Неверный формат файла. Загрузите CSV-файл."), 400

    try:
        # Парсинг CSV-файла
        file_data = file.read().decode('utf-8')

        # Убираем BOM, если он есть
        if file_data.startswith('\ufeff'):
            file_data = file_data[1:]

        csv_reader = csv.reader(file_data.splitlines(), delimiter=';')
        headers = next(csv_reader)
        records = []
        for row in csv_reader:
            records.append({headers[i]: value for i, value in enumerate(row)})
        
        # Возвращаем данные с порядком столбцов
        return jsonify({
            "success": True,
            "headers": headers,  # Важно: исходный порядок столбцов
            "records": records
        })
    except Exception as e:
        return jsonify(success=False, error=f"Ошибка обработки файла: {str(e)}"), 500

import os

@app.route('/save-my-template', methods=['POST'])
def save_my_template():
    data = request.get_json()

    if not data:
        return jsonify(success=False, error="Данные не переданы."), 400

    try:
        # Получаем данные изображения и имени файла
        image_data = data.get("imageData")
        filename = data.get("filename", "template.pdf")  # Имя PDF-файла

        if not image_data or not filename:
            return jsonify(success=False, error="Отсутствуют данные изображения или имя файла."), 400

        # Декодируем изображение из Base64
        image_data = image_data.split(",")[1]  # Убираем префикс 'data:image/png;base64,' 
        image_bytes = base64.b64decode(image_data)

        # Конвертируем в изображение
        image = Image.open(io.BytesIO(image_bytes))
        os.makedirs(os.getcwd(), exist_ok=True)  # Используем текущую директорию

        # Сохраняем PDF-файл
        pdf_output_path = os.path.join(os.getcwd(), filename)
        if image.mode == 'RGBA':  # Конвертируем, если изображение в формате RGBA
            image = image.convert('RGB')
        image.save(pdf_output_path, "PDF", resolution=100.0)

        # Создаем ZIP-архив с PDF-файлами
        zip_path = os.path.join(os.getcwd(), "templates.zip")
        with zipfile.ZipFile(zip_path, "w") as zipf:
            # Добавляем только PDF-файлы
            for file in os.listdir(os.getcwd()):
                file_path = os.path.join(os.getcwd(), file)
                if file.endswith(".pdf"):  # Добавляем только PDF
                    zipf.write(file_path, os.path.basename(file_path))

        # Возвращаем ссылку на скачивание архива
        return jsonify(success=True, downloadUrl=f"/download/{os.path.basename(zip_path)}")

    except Exception as e:
        logger.error(f"Ошибка при сохранении шаблона: {e}")
        return jsonify(success=False, error=f"Ошибка сохранения шаблона: {str(e)}"), 500

@app.route("/download/<path:filename>")
def download_file(filename):
    try:
        return send_file(filename, as_attachment=True)
    except Exception as e:
        logger.error(f"Ошибка при загрузке файла {filename}: {e}")
        return jsonify({"error": "Ошибка при загрузке файла."}), 500

# Работа с документами word
@app.route('/document')
def document():
    return render_template('document.html')

@app.route("/document/download", methods=["GET", "POST"])
def upload_files():
    if request.method == "POST":
        if "csvFile" not in request.files or "wordFile" not in request.files:
            return "Файлы не загружены", 400
        
        csv_file = request.files["csvFile"]
        word_file = request.files["wordFile"]
        
        try:
            # Чтение CSV с сохранением оригинальных названий
            df = pd.read_csv(csv_file, encoding='utf-8', delimiter=';', header=0)
            df.columns = [col.strip() for col in df.columns]  # Только пробелы
            
            logger.info(f"CSV columns: {df.columns.tolist()}")
            
            if df.empty:
                return "CSV файл пуст", 400

            # Сохраняем Word-файл
            original_word_path = os.path.join(temp_folder, "original.docx")
            word_file.save(original_word_path)
            
            # Анализ тегов
            doc = Document(original_word_path)
            used_tags = extract_tags(doc)
            logger.info(f"Найдены теги в документе: {used_tags}")
            
            # Определение используемых столбцов
            used_columns = [col for col in df.columns if col in used_tags]
            logger.info(f"Совпадающие столбцы: {used_columns}")
            
            if not used_columns:
                return "Нет совпадений между тегами документа и CSV", 400
                
            filename_column = used_columns[0]
            logger.info(f"Выбран столбец для имен файлов: {filename_column}")
            
            output_files = []
            generated_names = set()

            # Генерация документов
            for index, row in df.iterrows():
                # Формирование имени файла
                raw_name = str(row[filename_column])
                clean_name = re.sub(r'[\\/*?:"<>|]', '', raw_name).strip()
                safe_name = clean_name.replace(' ', '_')[:50]  # Ограничение длины
                
                # Уникализация имени
                counter = 1
                final_name = f"{safe_name}.docx"
                while final_name in generated_names:
                    final_name = f"{safe_name}_{counter}.docx"
                    counter += 1
                generated_names.add(final_name)
                
                # Копирование шаблона
                doc_path = os.path.join(temp_folder, final_name)
                copyfile(original_word_path, doc_path)
                
                # Замена тегов
                doc = Document(doc_path)
                replacements = {col: str(row[col]).strip() for col in df.columns}
                
                def replace_text(element):
                    if element.text:
                        for tag, value in replacements.items():
                            element.text = element.text.replace(tag, value)
                
                # Обработка всего документа
                for paragraph in doc.paragraphs:
                    replace_text(paragraph)
                    
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                replace_text(paragraph)
                
                doc.save(doc_path)
                output_files.append(doc_path)

            # Архивирование
            zip_path = os.path.join(temp_folder, "documents.zip")
            with zipfile.ZipFile(zip_path, "w") as zipf:
                for file in output_files:
                    zipf.write(file, os.path.basename(file))
            
            return send_file(zip_path, as_attachment=True)
        
        except Exception as e:
            logger.error(f"Ошибка: {str(e)}", exc_info=True)
            return f"Ошибка обработки: {str(e)}", 500
    
    return render_template("document.html")


@app.route("/output/<filename>")
def download_doc(filename):
    return send_from_directory('output', filename, as_attachment=True)

if __name__ == '__main__':
    app.run(debug=True)